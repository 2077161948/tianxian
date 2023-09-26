import io
import re

import pdfplumber

import fitz
import docx
import PyPDF2
from PyPDF2 import PdfReader

import PyPDF2
import fitz
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import PyPDF2
import fitz
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBox, LTTextLine, LTChar
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def extract_text_with_attributes_from_pdf(file_path):
    text_with_attributes = ""
    document = Document()
    parser = PDFParser(open(file_path, 'rb'))
    pdf = PDFDocument(parser)
    rsrcmgr = PDFResourceManager()
    laparams = LAParams()
    device = PDFPageAggregator(rsrcmgr, laparams=laparams)
    interpreter = PDFPageInterpreter(rsrcmgr, device)
    page_index = 0
    for page in PDFPage.create_pages(pdf):
        interpreter.process_page(page)
        layout = device.get_result()
        width = layout.width
        height = layout.height
        create_section(document, height, width, page_index)
        for element in layout:
            if isinstance(element, (LTTextBox, LTTextLine)):
                for text_line in element:
                    paragraph = document.add_paragraph()
                    for character in text_line:
                        if isinstance(character, LTChar):
                            text = character.get_text()
                            fontname = character.fontname if hasattr(character, 'fontname') else ""
                            fontsize = character.size if hasattr(character, 'size') else ""
                            color = character.graphicstate.ncolor
                            x0,x1,y0,y1 = character.x0,character.x1,character.y0,character.y1
                            print(text,fontname,fontsize)
                            #line_spacing = paragraph.paragraph_format
                            # run = paragraph.add_run(text)
                            # line_spacing.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 设置为固定行高
                            # line_spacing.line_spacing = Pt(fontsize)  # 设置行高为18磅
                            # run.font.size = Pt(fontsize)  # 替换为您想要的字体大小
                            # # 设置行高
                            # paragraph.paragraph_format.line_spacing = 1.5  # 替换为您想要的行高倍数
                            # run.font.name = fontname
                            # r, g, b = calculate_color(color)
                            # # 设置字体颜色
                            # run.font.color.rgb = RGBColor(r, g, b)  # 替换为您想要的RGB颜色值
                    save_docx(document)

    return text_with_attributes


def save_docx(document):
    # 保存Word文档
    document.save("custom_format.docx")


def paragraph_adds(paragraph, size, color, fontname, text):
    line_spacing = paragraph.paragraph_format
    run = paragraph.add_run(text)
    line_spacing.line_spacing_rule = WD_LINE_SPACING.EXACTLY  # 设置为固定行高
    line_spacing.line_spacing = Pt(size)  # 设置行高为18磅
    run.font.size = Pt(size)  # 替换为您想要的字体大小
    # 设置行高
    paragraph.paragraph_format.line_spacing = 1.5  # 替换为您想要的行高倍数
    run.font.name = fontname
    r,g,b = calculate_color(color)
    # 设置字体颜色
    run.font.color.rgb = RGBColor(r,g,b)  # 替换为您想要的RGB颜色值


def create_section(document, height, width, page):
    section = document.sections[page]
    if width > height:
        section.orientation = WD_ORIENTATION.LANDSCAPE
    # 设置自定义页面尺寸（可选）
    section.page_width = Pt(width)  # 设置页面宽度为792磅
    section.page_height = Pt(height)  # 设置页面高度为612磅


def calculate_color(color):
    R, G, B = color
    r = int(round(R * 255))
    g = int(round(G * 255))
    b = int(round(B * 255))
    #hex_color = '#{0:02X}{1:02X}{2:02X}'.format(r, g, b)
    return r,g,b


pdf_file = "C:/Users/Administrator/Desktop/经典算法大全_加水印.pdf"

# 调用函数提取图像的大小和位置信息
extract_text_with_attributes_from_pdf(pdf_file)


# doc = Document()
# for i in range(5):
#     if i != 3:
#         paragraph = doc.add_paragraph()
#         paragraph_adds(paragraph,size=60,color=(0.5,0,0),fontname='微软雅黑',text='qweeeee')
# save_docx(doc)